"""Modern Word (.docx) -> text + tables. Backend: python-docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document

from ..types import ParsedDoc, Table


def parse(path: str | Path) -> ParsedDoc:
    """Extract paragraphs as prose and tables as Table objects.

    Returns ParsedDoc with text=joined paragraphs and tables=list of Table.
    """
    path = Path(path)
    d = Document(path)

    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())

    tables: list[Table] = []
    for i, tbl in enumerate(d.tables):
        rows = [[cell.text for cell in row.cells] for row in tbl.rows]
        tables.append(Table(name=f"table_{i}", rows=rows))

    return ParsedDoc(
        text=text,
        tables=tables,
        metadata={"filename": path.name, "type": "docx"},
        source_ref={"filename": path.name},
    )
